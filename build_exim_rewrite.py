from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("/Users/saiyamchaplot/Documents/ChatGPT/Hackathon BWMI/EXIM_Preflight_Hackathon_Strategy_Rewritten.docx")

# compact_reference_guide preset, plus named visual overrides for this strategy brief.
COLORS = {
    "navy": "17324D",
    "blue": "2E74B5",
    "teal": "187D7A",
    "green": "2F7D57",
    "red": "A63838",
    "gold": "B87414",
    "ink": "2F3A45",
    "muted": "6B7280",
    "line": "D7E0E8",
    "blue_fill": "E8EEF5",
    "teal_fill": "E6F3F1",
    "green_fill": "E9F4EE",
    "red_fill": "F8E8E8",
    "gold_fill": "FFF3D9",
    "gray_fill": "F2F4F7",
    "white": "FFFFFF",
}

PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}
CALLOUT_MARGINS_DXA = {"top": 120, "bottom": 120, "start": 160, "end": 160}


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(run, *, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, margins):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margins[side]))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, *, color="D7E0E8", size=5, outer=True, inner=True):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        enabled = outer if edge in ("top", "left", "bottom", "right") else inner
        node.set(qn("w:val"), "single" if enabled else "nil")
        if enabled:
            node.set(qn("w:sz"), str(size))
            node.set(qn("w:space"), "0")
            node.set(qn("w:color"), color)


def set_cell_left_border(cell, color: str, size=18):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "right", "bottom"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "nil")
    left = borders.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        borders.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), "0")
    left.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, *, indent_dxa=TABLE_INDENT_DXA):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)


def set_paragraph_border_bottom(paragraph, color: str, size=8, space=4):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_paragraph_callout(paragraph, *, fill: str, accent: str):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")

    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), accent)
    p_bdr.append(left)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=COLORS["muted"])


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    r_pr.append(r_style)
    run.append(r_pr)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def configure_numbering(doc: Document):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    next_abstract = max(abstract_ids, default=-1) + 1
    next_num = max(num_ids, default=0) + 1

    def make_num(fmt: str, lvl_text: str):
        nonlocal next_abstract, next_num
        abstract_id = next_abstract
        num_id = next_num
        next_abstract += 1
        next_num += 1

        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text_el = OxmlElement("w:lvlText")
        lvl_text_el.set(qn("w:val"), lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "271")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.extend([tabs, ind, spacing])
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), "Calibri")
        r_fonts.set(qn("w:hAnsi"), "Calibri")
        r_pr.append(r_fonts)
        lvl.extend([start, num_fmt, lvl_text_el, suff, lvl_jc, p_pr, r_pr])
        abstract.append(lvl)
        numbering.append(abstract)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        numbering.append(num)
        return num_id

    return make_num("bullet", "•"), make_num("decimal", "%1.")


def apply_num(paragraph, num_id: int):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(COLORS["ink"])
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = rgb(COLORS["navy"])
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.line_spacing = 1.0

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    subtitle.font.size = Pt(13)
    subtitle.font.color.rgb = rgb(COLORS["muted"])
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle.paragraph_format.line_spacing = 1.15

    for style_name, size, color, before, after in (
        ("Heading 1", 16, COLORS["navy"], 18, 10),
        ("Heading 2", 13, COLORS["blue"], 14, 7),
        ("Heading 3", 12, COLORS["teal"], 10, 5),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    custom = {
        "Section Kicker": (10, COLORS["teal"], True, False, 0, 4, 1.0),
        "Lead": (13, COLORS["navy"], False, False, 0, 10, 1.2),
        "Small Muted": (9, COLORS["muted"], False, False, 0, 4, 1.1),
        "Table Citation": (9, COLORS["muted"], False, False, 4, 4, 1.1),
        "Step Title": (11.5, COLORS["navy"], True, False, 6, 2, 1.1),
        "Step Body": (10.5, COLORS["ink"], False, False, 0, 5, 1.2),
    }
    for name, (size, color, bold, italic, before, after, spacing) in custom.items():
        if name in styles:
            style = styles[name]
        else:
            style = styles.add_style(name, 1)
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = bold
        style.font.italic = italic
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = spacing


def setup_section(section, *, first=False):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    if first:
        section.different_first_page_header_footer = True


def configure_running_furniture(doc: Document):
    section = doc.sections[0]
    setup_section(section, first=True)
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("EXIM PRE-FLIGHT  |  HACKATHON BUILD STRATEGY")
    set_run_font(run, size=8.5, color=COLORS["navy"], bold=True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Build What Moves India  |  23 August 2026  |  ")
    set_run_font(run, size=8.5, color=COLORS["muted"])
    add_page_field(p)

    first_footer = section.first_page_footer
    p = first_footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Independent hackathon strategy  |  Prepared 23 August 2026")
    set_run_font(run, size=8.5, color=COLORS["muted"])


def add_kicker(doc, text):
    p = doc.add_paragraph(style="Section Kicker")
    p.paragraph_format.keep_with_next = True
    p.add_run(text.upper())
    return p


def add_section_heading(doc, number, kicker, title):
    add_kicker(doc, kicker)
    p = doc.add_paragraph(style="Heading 1")
    p.add_run(f"{number}. {title}")
    return p


def add_body(doc, text, *, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead, rest = text.split(bold_lead, 1)
        p.add_run(lead).bold = True
        p.add_run(rest)
    else:
        p.add_run(text)
    return p


def add_bullet(doc, text, bullet_num_id, *, bold_lead=None):
    p = doc.add_paragraph()
    apply_num(p, bullet_num_id)
    if bold_lead and text.startswith(bold_lead):
        lead, rest = text.split(bold_lead, 1)
        p.add_run(lead).bold = True
        p.add_run(rest)
    else:
        p.add_run(text)
    return p


def add_number(doc, text, decimal_num_id, *, bold_lead=None):
    p = doc.add_paragraph()
    apply_num(p, decimal_num_id)
    if bold_lead and text.startswith(bold_lead):
        lead, rest = text.split(bold_lead, 1)
        p.add_run(lead).bold = True
        p.add_run(rest)
    else:
        p.add_run(text)
    return p


def add_callout(doc, label, text, *, tone="blue"):
    palette = {
        "blue": (COLORS["blue"], COLORS["blue_fill"]),
        "teal": (COLORS["teal"], COLORS["teal_fill"]),
        "green": (COLORS["green"], COLORS["green_fill"]),
        "red": (COLORS["red"], COLORS["red_fill"]),
        "gold": (COLORS["gold"], COLORS["gold_fill"]),
        "navy": (COLORS["navy"], COLORS["blue_fill"]),
    }
    accent, fill = palette[tone]
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.10)
    p.paragraph_format.right_indent = Inches(0.06)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.keep_together = True
    set_paragraph_callout(p, fill=fill, accent=accent)
    label_run = p.add_run(label)
    set_run_font(label_run, size=11, color=accent, bold=True)
    label_run.add_break()
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=COLORS["ink"])
    return p


def add_table(doc, headers, rows, widths_dxa, *, header_fill=None, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, header_fill or COLORS["navy"])
        set_cell_margins(cell, CELL_MARGINS_DXA)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.1
        run = p.add_run(header)
        set_run_font(run, size=9.5, color=COLORS["white"], bold=True)
    set_repeat_table_header(table.rows[0])

    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        for col_idx, value in enumerate(values):
            cell = cells[col_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, CELL_MARGINS_DXA)
            if row_idx % 2 == 1:
                set_cell_shading(cell, COLORS["gray_fill"])
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(str(value))
            set_run_font(run, size=font_size, color=COLORS["ink"])
    set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA)
    set_table_borders(table, color=COLORS["line"], size=5)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    after.paragraph_format.line_spacing = 1
    return table


def add_step(doc, number, title, body, proof=None):
    p = doc.add_paragraph(style="Step Title")
    p.paragraph_format.keep_with_next = True
    marker = p.add_run(f"{number:02d}")
    set_run_font(marker, size=10, color=COLORS["teal"], bold=True)
    sep = p.add_run("  ")
    set_run_font(sep, size=10)
    run = p.add_run(title)
    set_run_font(run, size=11.5, color=COLORS["navy"], bold=True)
    p2 = doc.add_paragraph(style="Step Body")
    p2.paragraph_format.left_indent = Inches(0.32)
    p2.add_run(body)
    if proof:
        p3 = doc.add_paragraph(style="Small Muted")
        p3.paragraph_format.left_indent = Inches(0.32)
        label = p3.add_run("Proof: ")
        set_run_font(label, size=9, color=COLORS["teal"], bold=True)
        p3.add_run(proof)


def add_source(doc, title, url, note, decimal_num_id):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    apply_num(p, decimal_num_id)
    add_hyperlink(p, title, url)
    p2 = doc.add_paragraph(style="Small Muted")
    p2.paragraph_format.left_indent = Inches(0.22)
    p2.add_run(note)


def add_cover(doc):
    p = doc.add_paragraph(style="Section Kicker")
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run("HACKATHON BUILD STRATEGY")

    p = doc.add_paragraph(style="Title")
    p.add_run("EXIM Pre-Flight\nCompliance Engine")
    set_paragraph_border_bottom(p, COLORS["blue"], size=8, space=5)

    p = doc.add_paragraph(style="Subtitle")
    p.add_run("A plain-English build brief for Build What Moves India")

    add_callout(
        doc,
        "Core decision",
        "Pitch the long-term platform. Build and demo one narrow import journey that works from start to finish.",
        tone="navy",
    )

    p = doc.add_paragraph(style="Lead")
    p.add_run(
        "The prototype should help one small Indian importer decide whether one proposed shipment is ready to move, what is still missing, what customs may cost, and what to do next."
    )

    add_table(
        doc,
        ["Direction", "Primary user", "Demo product", "Route"],
        [["Import into India", "Small MSME importer", "100 Bluetooth speakers", "China to JNPT"]],
        [1900, 2500, 2800, 2160],
        font_size=9.3,
    )

    add_callout(
        doc,
        "Document status",
        "Submission rules were checked against the official builder brief and FAQ on 23 August 2026. Recheck the official pages immediately before submission because requirements can change.",
        tone="gold",
    )


def build_document():
    doc = Document()
    doc.core_properties.title = "EXIM Pre-Flight Compliance Engine - Hackathon Build Strategy"
    doc.core_properties.subject = "Build What Moves India strategy, demo journey, architecture, and submission plan"
    doc.core_properties.author = ""
    doc.core_properties.keywords = "EXIM, import compliance, hackathon, OpenAI, Codex"
    configure_styles(doc)
    configure_running_furniture(doc)
    bullet_num_id, decimal_num_id = configure_numbering(doc)

    add_cover(doc)
    doc.add_page_break()

    add_section_heading(doc, 1, "Decision", "What to Build")
    add_callout(
        doc,
        "Recommendation",
        "Build a source-linked import readiness checker for one fictional shipment: 100 Bluetooth speakers from China to Jawaharlal Nehru Port (JNPT), also known as Nhava Sheva.",
        tone="teal",
    )
    add_body(
        doc,
        "The long-term idea is a pre-flight compliance operating system for Indian trade. That idea is too broad for a short hackathon. The submission should prove one complete journey that reviewers can use and test."
    )
    doc.add_paragraph("The product promise", style="Heading 2")
    add_callout(
        doc,
        "The user question",
        "Before I pay the supplier or ship the goods, can this shipment move through Indian customs, what is missing, what may it cost, and what must I do next?",
        tone="navy",
    )
    doc.add_paragraph("The success test", style="Heading 2")
    for item in (
        "A reviewer can complete the main journey from a fresh browser session.",
        "The result changes when the reviewer fixes a real document mismatch.",
        "Every compliance outcome shows its source, rule date, assumptions, and confidence.",
        "Every simulated government response is labelled as synthetic.",
        "The same structured input produces the same rule result.",
    ):
        add_bullet(doc, item, bullet_num_id)
    add_callout(
        doc,
        "The hard scope rule",
        "Do not add a feature unless it makes this one journey clearer, more credible, or more complete.",
        tone="gold",
    )

    add_section_heading(doc, 2, "Problem and positioning", "Why This Problem Is Worth Solving")
    add_body(
        doc,
        "A small importer can discover a missing approval, weak invoice description, or certificate mismatch only after goods have been ordered or shipped. At that point, inventory and working capital may be trapped while the business resolves the issue."
    )
    doc.add_paragraph("Why the idea fits the hackathon", style="Heading 2")
    for lead, rest in (
        ("One clear user.", " A small or first-time Indian importer with limited customs knowledge."),
        ("One costly decision.", " Decide whether to pay or ship before the compliance picture is clear."),
        ("A complete journey.", " Collect facts, ask focused questions, identify blockers, validate evidence, and produce the next steps."),
        ("Meaningful AI use.", " Read unstructured invoices and certificates, ask for missing product facts, compare evidence, and explain rule outcomes."),
        ("Honest simulation.", " Use fictional identities and mocked government responses while keeping the core product logic real."),
    ):
        add_bullet(doc, lead + rest, bullet_num_id, bold_lead=lead)
    doc.add_paragraph("The evidence gap", style="Heading 2")
    add_callout(
        doc,
        "Close this before submission",
        "The official brief asks builders to choose a real problem they have faced. The strategy still needs one precise, truthful incident from the team or a directly interviewed importer, customs broker, freight forwarder, or import-export adviser. Do not invent a founder story.",
        tone="red",
    )
    for item in (
        "What product was involved?",
        "What requirement or document mismatch was missed?",
        "When was the problem discovered?",
        "Which services, portals, or professionals were involved?",
        "What time, money, or operational delay did it cause?",
        "What did the business do to resolve it?",
    ):
        add_bullet(doc, item, bullet_num_id)

    add_section_heading(doc, 3, "Prototype scope", "The Narrow Demonstration")
    add_table(
        doc,
        ["Field", "Hackathon choice"],
        [
            ["Direction", "Import into India only"],
            ["User", "Small Indian MSME importer"],
            ["Product", "100 Bluetooth speakers"],
            ["Origin", "China"],
            ["Destination port", "JNPT / Nhava Sheva"],
            ["Identity", "Fictional importer and fictional Importer Exporter Code (IEC)"],
            ["Regulatory coverage", "Only rules that the team has verified and date-stamped"],
        ],
        [2500, 6860],
    )
    doc.add_paragraph("What the prototype must prove", style="Heading 2")
    for item in (
        "Turn an invoice and product specification into one structured shipment record.",
        "Ask only questions that can change classification or compliance.",
        "Run a curated rule set with traceable, repeatable logic.",
        "Calculate duties from explicit, date-stamped rate fixtures.",
        "Check whether uploaded evidence matches the shipment record.",
        "Turn each blocker into an ordered task with a clear owner and destination service.",
    ):
        add_bullet(doc, item, bullet_num_id)
    add_callout(
        doc,
        "Regulatory caution",
        "Do not claim that every Bluetooth speaker follows the same compliance path. Product design, radio features, battery, model, manufacturer, intended use, and current notifications can change the answer. Ask the missing questions and show the supported rule date.",
        tone="red",
    )

    add_section_heading(doc, 4, "Reviewer journey", "The Seven-Step Demo")
    add_step(
        doc,
        1,
        "Start a pre-flight check",
        "Enter Import, Bluetooth speaker, China, JNPT, 100 units, invoice value and currency, freight and insurance or Incoterm, and the proposed shipment date. Upload a synthetic invoice and specification sheet.",
        "The shipment can be created without using real personal or government data.",
    )
    add_step(
        doc,
        2,
        "Extract the product facts",
        "OpenAI reads the files and proposes the category, brand, model, radio features, power source, battery, manufacturer, origin, price, and quantity.",
        "The reviewer sees the extracted fields and can correct them before rules run.",
    )
    add_step(
        doc,
        3,
        "Ask the few questions that matter",
        "Ask whether the product also has Wi-Fi, whether it is a smart speaker, whether the battery is built in, and whether the enclosure contains one or more speaker units. Ask only questions that can change the result.",
        "At least one answer changes a candidate classification or requirement.",
    )
    add_step(
        doc,
        4,
        "Show business readiness and the first result",
        "Use a fictional importer such as Aarav Audio Traders LLP and a fictional IEC such as DEMO0123456. Show IEC active, Goods and Services Tax Identification Number (GSTIN) available, and ICEGATE or port bank-code readiness not yet verified.",
        "The screen is visibly labelled: synthetic demonstration data; no live government system was accessed.",
    )
    add_step(
        doc,
        5,
        "Explain why the shipment is blocked",
        "Show a clear state such as DO NOT SHIP YET, followed by proposed Indian Trade Classification (ITC-HS), candidate alternatives, supported approvals, missing evidence, customs estimate, assumptions, sources, and effective dates.",
        "Each blocker has a reason, authority or service, owner, required file, and blocking status.",
    )
    add_step(
        doc,
        6,
        "Upload evidence and catch a deliberate mismatch",
        "Upload synthetic certificates and an amended invoice. Compare model, manufacturer, validity, product category, quantity, and evidence coverage. Include one controlled error: certificate model BS-100 while the invoice lists BS-100X.",
        "The engine identifies the mismatch and explains who must correct it.",
    )
    add_step(
        doc,
        7,
        "Fix, re-run, and generate the packet",
        "After the mismatch is corrected, re-run the same rules. Show READY FOR CUSTOMS PREPARATION only when no supported blocker remains. Generate a shareable Shipment Compliance Packet with the summary, proposed classification, requirements, duty estimate, evidence checklist, assumptions, sources, rule version, and timestamp.",
        "The result changes from blocked to ready within the supported rule set; the product never guarantees customs clearance.",
    )

    add_section_heading(doc, 5, "Implementation honesty", "What Is Real and What Is Simulated")
    add_body(
        doc,
        "Build the parts that prove product value. Simulate the parts that require government access, private data, credentials, payments, or production permissions."
    )
    add_table(
        doc,
        ["Build for real", "Simulate and label"],
        [
            ["Invoice and specification extraction", "IEC status response"],
            ["Focused clarification questions", "Bureau of Indian Standards (BIS) status"],
            ["Structured shipment record", "Wireless Planning and Coordination (WPC) application status"],
            ["Rule evaluation for the supported case", "ICEGATE registration or filing"],
            ["Duty arithmetic from curated fixtures", "Government payment"],
            ["Document consistency checks", "Portal submission"],
            ["Action plan and compliance packet", "Customs-clearance outcome"],
            ["Source links, rule dates, and assumptions", "Any live identity or government account"],
        ],
        [4680, 4680],
    )
    add_callout(
        doc,
        "Required disclosure",
        "This is an independent hackathon prototype. It uses fictional identities and synthetic government responses. No live government system was accessed, and the result is not a legal opinion or a customs-clearance guarantee.",
        tone="gold",
    )
    doc.add_paragraph("Plan for future connectors without depending on them", style="Heading 2")
    add_body(
        doc,
        "Put every government integration behind a small provider interface. During the hackathon, use a synthetic provider. Later, an approved connector can implement the same interface. The main demo must still work when no live endpoint exists."
    )

    add_section_heading(doc, 6, "Technical design", "How OpenAI and the Rules Engine Share the Work")
    add_callout(
        doc,
        "System rule",
        "AI may read, extract, compare, organise, and explain. It may not invent a classification, duty rate, approval requirement, or final compliance conclusion.",
        tone="red",
    )
    add_table(
        doc,
        ["Layer", "Responsibility"],
        [
            ["OpenAI extraction", "Turn descriptions, invoices, specifications, and certificates into structured fields."],
            ["Clarification planner", "Find only the unanswered product facts that can change the result."],
            ["Rules engine", "Apply the supported, date-effective rules to the structured shipment record."],
            ["Duty calculator", "Use explicit formulas and rate fixtures; return inputs, steps, and assumptions."],
            ["Evidence validator", "Compare document facts with the shipment and the evidence each rule requires."],
            ["OpenAI explanation", "Translate deterministic outcomes into plain-English reasons and next actions."],
        ],
        [2600, 6760],
    )
    doc.add_paragraph("The single source of truth", style="Heading 2")
    add_body(
        doc,
        "Store one persistent shipment record containing the parties, product attributes, candidate ITC-HS codes, route, commercial terms, applicable rules, documents, registrations, duties, risks, assumptions, actions, source versions, and timestamps. Every model call, rule, calculation, and screen reads from or writes to this record."
    )
    doc.add_paragraph("Required controls", style="Heading 2")
    for item in (
        "Date-effective rule storage and explicit rule versions.",
        "Clear thresholds and blocking conditions.",
        "Traceable duty formulas and rate fixtures.",
        "Official source links on every supported outcome.",
        "Regression fixtures for each supported scenario.",
        "Stable results for identical structured inputs.",
        "Human escalation when sources conflict or required facts remain unknown.",
    ):
        add_bullet(doc, item, bullet_num_id)
    add_callout(
        doc,
        "Credible Codex statement",
        "Codex helped design and implement the rule pipeline, create supported shipment fixtures, build the document-matching flow, and write regression tests that check identical inputs produce identical compliance results.",
        tone="blue",
    )

    add_section_heading(doc, 7, "Scope control", "What Not to Build")
    add_body(doc, "Exclude work that does not strengthen the single importer journey.")
    add_table(
        doc,
        ["Exclude", "Reason"],
        [
            ["Export compliance", "Adds destination-country rules, incentives, certificates, and remittance workflows."],
            ["Every product or ITC-HS code", "Creates a broad but unreliable classifier."],
            ["Every regulator", "Makes rule coverage shallow and hard to verify."],
            ["Automatic portal login or filing", "Adds credentials, one-time passwords, CAPTCHA, signatures, portal changes, and legal responsibility."],
            ["Real IEC, PAN, GST, or bank data", "Creates unnecessary privacy and security risk."],
            ["Logistics booking or payments", "Distracts from the compliance decision."],
            ["Broker or consultant marketplace", "Adds a second business model before the core product works."],
            ["Complex administration dashboard", "Reviewers will test the citizen journey."],
            ["Generic chat interface", "Hides the decision structure and looks like a thin wrapper."],
            ["Customs-clearance guarantee", "Cannot be defended legally or operationally."],
        ],
        [3300, 6060],
        font_size=9.2,
    )
    add_callout(
        doc,
        "Scope test",
        "One complete journey will score better than ten features that work only halfway.",
        tone="navy",
    )

    add_section_heading(doc, 8, "Judging and proof", "How the Narrow Build Competes")
    add_table(
        doc,
        ["Judging area", "What the submission must prove"],
        [
            ["Problem", "One truthful incident shows why late discovery of a missing requirement matters."],
            ["Working build", "A reviewer completes the main journey and sees the result change after fixing evidence."],
            ["Usability", "A structured mobile-friendly flow explains terms, assumptions, blockers, and next actions."],
            ["Product thinking", "The scope is narrow, the questions are purposeful, and unsupported cases escalate."],
            ["End-to-end thinking", "The build includes the data model, rules, calculations, validation, and process owners."],
            ["Honesty", "Synthetic data, mocked integrations, limitations, and no-clearance guarantee are explicit."],
        ],
        [2400, 6960],
    )
    add_callout(
        doc,
        "Competitive view",
        "The concept can compete if the build is narrow, polished, and testable. The biggest remaining risk is not the architecture; it is weak proof that the team understands a real importer incident firsthand.",
        tone="teal",
    )

    add_section_heading(doc, 9, "Submission narrative", "Two-Minute Demonstration Video")
    add_table(
        doc,
        ["Time", "Purpose", "Show"],
        [
            ["0:00-0:08", "Problem", "A small importer can discover missing approvals only after ordering or shipping."],
            ["0:08-0:45", "Citizen journey", "Enter the shipment, upload the invoice, review extracted facts, answer focused questions, and reveal blockers plus the duty estimate."],
            ["0:45-1:00", "Resolution", "Upload synthetic evidence, catch one model mismatch, correct it, and generate the packet."],
            ["1:00-1:25", "How it works", "Explain OpenAI extraction, structured shipment record, rules engine, calculator, validator, and source tracking."],
            ["1:25-1:45", "Honesty", "State the supported scope, synthetic identity, mocked government responses, and no-clearance guarantee."],
            ["1:45-2:00", "Scale", "Show how new products, rules, and approved connectors can use the same journey later."],
        ],
        [1400, 1900, 6060],
        font_size=9.0,
    )
    add_callout(
        doc,
        "Video discipline",
        "Use the first minute for the working citizen journey and the second minute for how you built it and why. Do not open with slides, market size, or a long monologue.",
        tone="blue",
    )

    add_section_heading(doc, 10, "Submission gate", "Rules to Confirm Before You Submit")
    add_body(
        doc,
        "The following requirements were confirmed on the official builder brief and FAQ on 23 August 2026. Check those pages again immediately before submission."
    )
    for item in (
        "Submit by 28 August 2026 at 8:00 PM IST; the official brief says there is no grace period.",
        "Provide a public browser link that opens without requesting access.",
        "If login is required, provide fictional consumer credentials; never provide a real password.",
        "Submit one video no longer than two minutes.",
        "Submit a project summary under 250 words.",
        "Participate solo or in a team of two. Both teammates must register separately and cross-reference registered email addresses.",
        "Build the main journey as a working prototype; a static design is not enough.",
        "Use Codex meaningfully and explain its contribution.",
        "Do not connect to a live government system unless organisers provide an approved public sandbox.",
        "Do not use real sensitive data or present the prototype as official or endorsed.",
    ):
        add_bullet(doc, item, bullet_num_id)

    add_section_heading(doc, 11, "Execution plan", "Six-Day Build Schedule")
    add_table(
        doc,
        ["Date", "Required output"],
        [
            ["23 Aug", "Lock the shipment scenario, supported inputs and outputs, verified rule scope, synthetic files, schemas, and test cases."],
            ["24 Aug", "Finish extraction, clarification logic, rule evaluation, duty calculation, mock connectors, and at least five automated scenarios."],
            ["25 Aug", "Finish the responsive journey, clear blocked/unknown/ready states, action plan, sources, assumptions, and rule dates."],
            ["26 Aug", "Finish evidence upload, model/manufacturer/expiry checks, packet generation, and loading/error/incomplete-data states."],
            ["27 Aug", "Test desktop, mobile, slow connection, and fresh-session flows; remove incomplete features; deploy; record video; draft summary."],
            ["28 Aug", "Test every link in an incognito browser, run the demo with a fresh fictional account, verify permissions, recheck official rules, and submit early."],
        ],
        [1500, 7860],
        font_size=9.1,
    )
    doc.add_paragraph("Minimum regression scenarios", style="Heading 2")
    for item in (
        "Happy path: every required item of evidence matches.",
        "Missing evidence: one required certificate is absent.",
        "Model mismatch: the certificate model differs from the invoice model.",
        "Classification uncertainty: a missing product function leaves two candidate codes.",
        "Commercial-data gap: freight or Incoterm is missing, so a complete estimate cannot be produced.",
    ):
        add_bullet(doc, item, bullet_num_id)
    add_callout(
        doc,
        "Final release gate",
        "Open the public build in an incognito window on desktop and mobile. Complete the journey from a fresh fictional account. Confirm that every link, upload, calculation, state change, source, and mock label works without help from the builder.",
        tone="green",
    )

    add_section_heading(doc, 12, "Final position", "The Pitch to Use")
    add_callout(
        doc,
        "Long-term vision",
        "A pre-flight compliance operating system for Indian trade.",
        tone="blue",
    )
    add_callout(
        doc,
        "Hackathon submission",
        "An independent, source-linked pre-flight checker that helps a small Indian importer understand whether one proposed Bluetooth-speaker shipment is ready for customs preparation, which approvals and documents are still missing, which services must be used, and what the estimated customs cost may be before the goods move.",
        tone="teal",
    )
    doc.add_paragraph("Defensible difference", style="Heading 2")
    add_body(
        doc,
        "The product does not replace the Directorate General of Foreign Trade (DGFT), ICEGATE, API Setu, BIS, WPC, or customs brokers. It keeps one structured record of the shipment, asks for missing product facts, applies a verified rule set, checks uploaded evidence, orders the next actions, and shows who owns each blocker."
    )
    add_callout(
        doc,
        "Demo line",
        "Know what can block the shipment before the shipment moves.",
        tone="navy",
    )
    doc.add_paragraph("Required limitation statement", style="Heading 2")
    add_body(
        doc,
        "This prototype supports one defined import scenario and a date-stamped rule set. It identifies unresolved pre-flight risks and produces an action plan. It does not provide a legal opinion or guarantee customs clearance."
    )

    add_section_heading(doc, 13, "Reference", "Plain-English Glossary")
    add_table(
        doc,
        ["Term", "Meaning in this document"],
        [
            ["AD code", "Authorized Dealer code used to link a port registration to the importer's bank branch."],
            ["API Setu", "A Government of India platform for approved application programming interface access."],
            ["BIS", "Bureau of Indian Standards."],
            ["CHA", "Customs House Agent; commonly called a customs broker."],
            ["DGFT", "Directorate General of Foreign Trade."],
            ["ETA", "Equipment Type Approval for supported wireless equipment cases."],
            ["GSTIN", "Goods and Services Tax Identification Number."],
            ["IEC", "Importer Exporter Code."],
            ["ICEGATE", "Indian Customs Electronic Gateway, the customs electronic service platform."],
            ["ITC-HS", "India's trade-classification code based on the Harmonized System."],
            ["JNPT", "Jawaharlal Nehru Port, also known as Nhava Sheva."],
            ["MSME", "Micro, small, or medium enterprise."],
            ["WPC", "Wireless Planning and Coordination Wing of the Department of Telecommunications."],
        ],
        [1900, 7460],
        font_size=9.0,
    )

    add_section_heading(doc, 14, "Sources", "Official Pages Used")
    add_body(
        doc,
        "These pages support the hackathon rules, the future connector concept, and the wireless-equipment service reference. Rules can change; recheck the pages before the final submission."
    )
    add_source(
        doc,
        "Build What Moves India - Builder brief",
        "https://buildwhatmovesindia.com/brief",
        "Challenge, working journey, safety rules, submission requirements, deadline, and judging criteria. Verified 23 August 2026.",
        decimal_num_id,
    )
    add_source(
        doc,
        "Build What Moves India - Frequently asked questions",
        "https://buildwhatmovesindia.com/faq",
        "Team size, Codex requirement, live-system restrictions, synthetic data, video, and submission details. Verified 23 August 2026.",
        decimal_num_id,
    )
    add_source(
        doc,
        "API Setu - IEC API collection",
        "https://directory.apisetu.gov.in/api-collection/iec#collection01",
        "Reference for a future approved IEC connector; not a required live dependency for the prototype.",
        decimal_num_id,
    )
    add_source(
        doc,
        "API Setu - Standard Operating Procedure",
        "https://apisetu.gov.in/sop",
        "Reference for approved API access and onboarding.",
        decimal_num_id,
    )
    add_source(
        doc,
        "Department of Telecommunications - Equipment Type Approval",
        "https://www.eservices.dot.gov.in/equipment-type-approval-eta",
        "Official wireless-equipment service reference. The prototype must still verify the product-specific applicability and effective rule date.",
        decimal_num_id,
    )
    add_callout(
        doc,
        "Source discipline",
        "Every compliance outcome in the prototype should show an official source, publication or effective date, rule version, assumptions, and confidence. When the verified rules do not support a clear answer, escalate instead of guessing.",
        tone="gold",
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    run = p.add_run("End of document")
    set_run_font(run, size=9, color=COLORS["muted"], italic=True)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
